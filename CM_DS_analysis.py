

#VERSION 5

import pandas as pd
import numpy as np
from docx import Document
from datetime import date
import glob, os

# Returns the current local date
today = date.today()

#User input
inputdir =os.getcwd()

for doc_file in glob.glob(os.path.join(inputdir,"*.doc*")):
    document = Document(doc_file) # need to make it auto input
lens=len(document.tables)
print(lens,"tables extracted")
df=pd.DataFrame()
dataframe_list=[]
# Extraction of tables from word file
for x in range(lens):
    table=document.tables[x]
    data = [[cell.text for cell in row.cells] for row in table.rows]
    df1 = pd.DataFrame(data)
    dataframe_list.append(df1)
    
dataframe_list2 = [x for x 
                   in dataframe_list 
                   if x.apply(lambda x: x.str.contains('Unique Identifier')).any().any()]

df_new2=pd.DataFrame()
dataframe_list3=[]
for x in (dataframe_list2):
    df_new1=pd.DataFrame(x)
    n=len(x.columns)

    if(n==8):
        df_new1.insert(3, "Spec_guar", "") 
        col_names=[0,1,2,3,4,5,6,7,8]
        df_new1.columns=col_names
        dataframe_list3.append(df_new1)     
        #print('Test')
        #print(df_new1)
    if(n<8):
        print("Error in table: Inconsistency with columns \n",x,"\n Please copy the above UID and check the datasheet and rerun the script. Press enter to continue")
        val=input()
        #print(val)
        os._exit(0)
    if(n==9): 
        #print('Else')
        dataframe_list3.append(df_new1)
        #print(df_new1)
df_new3=pd.DataFrame()

for x in (dataframe_list3):
    df_new2=pd.DataFrame(x)
    #print('Final_df')
    #print(df_new2)
    frames = [df_new3,df_new2]
    df_new3 = pd.concat(frames, ignore_index=True)

#print (dataframe_list3)
#print(df_new3)

column_names=['DS_Symbol','DS_Parameter','DS_Conditions','DS_Spec_guar','DS_Min','DS_Typ','DS_Max','DS_Unit',
              'Unique Identifier']
df_new3.columns=column_names
df_new13=  df_new3[df_new3["DS_Symbol"].str.contains("Symbol") == False]


column_names2=['CM_Symbol','CM_Parameter','CM_Conditions','CM_Spec_guar','CM_Min','CM_Typ','CM_Max','CM_Unit',
              'Unique Identifier','Datasheet']



# Remove spaces from cells
def trim_all_cols(df):
    trim_strings=lambda x:x.strip() if isinstance(x,str) else x
    return df.applymap(trim_strings)

df_new13=trim_all_cols(df_new13)

#df_new13.to_excel ("DS.xlsx", index = None, header=True) ## set current directory and adjust


#CM
for cm_file in glob.glob(os.path.join(inputdir,"*.xlsm*")):
    df2=pd.read_excel(open(cm_file, 'rb'),
              sheet_name='Parametric')  
new_header = df2.iloc[0] #grab the first row for the header
df2 = df2[1:] #take the data less the header row
df2.columns = new_header #set the header row as the df header
print(new_header)
df2 = df2[["Symbol","Parameter","Conditions","Specification Guaranteed by","Min","Typ","Max","Unit","Unique Identifier", "Datasheet"]]
df2.columns=column_names2
#df2.to_excel ('CM.xlsx', index = None, header=True) ## set current directory and adjust
#print(df2)

df3 = pd.merge(df_new13, df2, how='outer', on=['Unique Identifier'])


#data = df3.copy()                        


#data= trim_all_cols(data)
#data = data.astype({'Min': str, 'DS_Min': str,'Max': str, 'DS_Max': str,'Typ': str, 'DS_Typ': str}) 

df3['CM_Min']=df3['CM_Min'].apply(str) 
df3['CM_Max']=df3['CM_Max'].apply(str)
df3['CM_Typ']=df3['CM_Typ'].apply(str)
df3['CM_Symbol']=df3['CM_Symbol'].apply(str)
df3['CM_Parameter']=df3['CM_Parameter'].apply(str)
df3['CM_Conditions']=df3['CM_Conditions'].apply(str)
df3['CM_Spec_guar']=df3['CM_Spec_guar'].apply(str)
df3['Unique Identifier']=df3['Unique Identifier'].apply(str)
df3['CM_Unit']=df3['CM_Unit'].apply(str)
df3_Uni=df3["Unique Identifier"]
df3["Datasheet"]=df3["Datasheet"].apply(str)
df3_Uni=pd.Series(df3_Uni)


#Conversion to array for vectorization
DS_Min_array=df3["DS_Min"].to_numpy()
DS_Typ_array=df3["DS_Typ"].to_numpy()
DS_Max_array=df3["DS_Max"].to_numpy()
DS_Symbol_array=df3["DS_Symbol"].to_numpy()
DS_Parameter_array=df3["DS_Parameter"].to_numpy()
DS_Conditions_array=df3["DS_Conditions"].to_numpy()
DS_Spec_guar_array=df3["DS_Spec_guar"].to_numpy()
DS_Unit_array=df3["DS_Unit"].to_numpy()
DS_Datasheet_array=df3["Datasheet"].to_numpy()
df3_Uni_array=df3_Uni.to_numpy()


#String datatype for all dataframes 
DS_Max_array = DS_Max_array.astype('str')
DS_Min_array = DS_Min_array.astype('str')
DS_Typ_array = DS_Typ_array.astype('str')
DS_Symbol_array = DS_Symbol_array.astype('str')
DS_Parameter_array = DS_Parameter_array.astype('str')
DS_Conditions_array = DS_Conditions_array.astype('str')
DS_Spec_guar_array = DS_Spec_guar_array.astype('str')
DS_Unit_array = DS_Unit_array.astype('str')
df3_Uni_array = df3_Uni_array.astype('str')
DS_Datasheet_array = DS_Datasheet_array.astype('str')

#Handling missing values
DS_Typ_array[DS_Typ_array==''] = 'nan'
DS_Max_array[DS_Max_array==''] = 'nan'
DS_Min_array[DS_Min_array==''] = 'nan'
DS_Symbol_array[DS_Symbol_array==''] = 'nan'
DS_Parameter_array[DS_Parameter_array==''] = 'nan'
DS_Conditions_array[DS_Conditions_array==''] = 'nan'
DS_Spec_guar_array[DS_Spec_guar_array==''] = 'nan'
DS_Unit_array[DS_Unit_array==''] = 'nan'
df3_Uni_array[df3_Uni_array==''] = 'nan'
DS_Datasheet_array[DS_Datasheet_array==''] = 'nan'


df3['DS_Min']=DS_Min_array
df3['DS_Max']=DS_Max_array
df3['DS_Typ']=DS_Typ_array
df3['DS_Symbol']=DS_Symbol_array
df3['DS_Parameter']=DS_Parameter_array
df3['DS_Conditions']=DS_Conditions_array
df3['DS_Spec_guar']=DS_Spec_guar_array
df3['DS_Unit']=DS_Unit_array
df3['Datasheet']=DS_Datasheet_array
df3.rename(columns={"Datasheet":"Datasheet"}, inplace = True)


df3['Symbol OK'] = df3['DS_Symbol'] == df3['CM_Symbol']
df3['Parameter OK'] = df3['CM_Parameter'] == df3['DS_Parameter']
df3['Conditions OK'] = df3['CM_Conditions'] == df3['DS_Conditions']
df3['Specifications OK'] = df3['CM_Spec_guar'] == df3['DS_Spec_guar']

df3['Min OK'] = df3['CM_Min'] == df3['DS_Min'] 
df3['Typ OK'] = df3['CM_Typ'] == df3['DS_Typ']
df3['Max OK'] = df3['CM_Max'] == df3['DS_Max']

df3['Unit OK'] = df3['DS_Unit'] == df3['CM_Unit']

#Error for UIDs not included in Datasheet
df3["Datasheet_error"] = df3[["Datasheet", "DS_Min","DS_Max","DS_Typ"]].apply(
    lambda x: "Unique Identifier should be there in Datasheet"
    if (x.Datasheet == "Yes" and x.DS_Min == "nan" and x.DS_Max == "nan" and x.DS_Typ == "nan")
    else "No action needed",
    axis=1
)



# NDS + "" UIDs identifier
NDS=(np.char.startswith(df3_Uni_array, 'NDS'))
df3['NDS_identifier']=np.char.startswith(df3_Uni_array, 'NDS')

# remove nan in final result
df3=  df3[df3["Unique Identifier"].str.contains("nan") == False]

#COLOR_Highlighted
def highlight_cells(val, color_if_true, color_if_false):
    
    if bool(val == False) : color = color_if_true
    elif (val == 'Unique Identifier should be there in Datasheet') : color = color_if_true
    else : color = color_if_false
    return 'background-color: {}'.format(color)

output_file= inputdir + "\\"+str(today)+"-New-CM_VS_DS.xlsx"
Unstyled= inputdir + "\\"+str(today)+"Uncolored-New-CM_VS_DS.xlsx"

column_names3=['DS-Symbol','DS-Parameter','DS-Conditions','DS-Spec-guar','DS-Min','DS-Typ','DS-Max','DS-Unit',
              'Unique Identifier','CM-Symbol','CM-Parameter','CM-Conditions','CM-Spec-guar','CM-Min','CM-Typ','CM-Max','CM-Unit',
              'Datasheet','Symbol OK','Parameter OK','Conditions OK','Specifications OK','Min OK', 'Typ OK','Max OK','Unit OK','Datasheet-error','NDS-Identifier']
df3.columns=column_names3
df3.style.applymap(highlight_cells, color_if_true='yellow', color_if_false='white' ,subset=['Symbol OK','Parameter OK','Conditions OK','Specifications OK','Min OK', 'Typ OK','Max OK','Unit OK','Datasheet-error'])\
        .to_excel(output_file, engine='openpyxl')


# Excel output of clean DS extracted from pdf without color
#df3.to_excel (Unstyled, index = None, header=True) ## set current directory and adjust

